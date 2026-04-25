"""导出服务：生成真实可打开的 DOCX / EPUB / PDF / Markdown。"""
from __future__ import annotations

import html
import io
import logging
import os
import re
import tempfile
from pathlib import Path
from typing import Iterator, List, Optional, Tuple

from domain.novel.repositories.novel_repository import NovelRepository
from domain.novel.repositories.chapter_repository import ChapterRepository
from domain.novel.entities.novel import Novel
from domain.novel.entities.chapter import Chapter
from domain.novel.value_objects.novel_id import NovelId
from domain.novel.value_objects.chapter_id import ChapterId

logger = logging.getLogger(__name__)


def _safe_filename_stem(title: str, max_len: int = 80) -> str:
    t = (title or "novel").strip()
    t = re.sub(r'[<>:"/\\|?*\x00-\x1f]', "_", t)
    t = t.replace(" ", "_").strip("._") or "novel"
    if len(t) > max_len:
        t = t[:max_len]
    return t


def _novel_id_str(novel: Novel) -> str:
    nid = novel.id
    return nid.value if hasattr(nid, "value") else str(nid)


def _chapter_display_title(ch: Chapter) -> str:
    prefix = f"第{ch.number}章"
    if ch.title and str(ch.title).strip():
        title = str(ch.title).strip()
        compact = re.sub(r"\s+", "", title)
        if re.match(r"^第[0-9一二三四五六七八九十百千零两]+章", compact):
            return title
        return f"{prefix}：{title}"
    return prefix


def _content_to_html_paragraphs(text: str) -> str:
    raw = (text or "").replace("\r\n", "\n").replace("\r", "\n")
    parts: List[str] = []
    for block in raw.split("\n"):
        line = block.strip()
        if line:
            parts.append(f"<p>{html.escape(line)}</p>")
    if not parts:
        return "<p></p>"
    return "\n".join(parts)


def _cjk_font_paths() -> Iterator[Path]:
    env = os.environ.get("PLOTPILOT_EXPORT_CJK_FONT", "").strip()
    if env:
        yield Path(env)
    if os.name == "nt":
        windir = os.environ.get("WINDIR", r"C:\Windows")
        fonts = Path(windir) / "Fonts"
        for name in (
            "msyh.ttc",
            "msyh.ttf",
            "simhei.ttf",
            "simsun.ttc",
            "simkai.ttf",
            "simfang.ttf",
            "simli.ttf",
        ):
            yield fonts / name
    else:
        for p in (
            "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
            "/usr/share/fonts/opentype/noto/NotoSansCJKsc-Regular.otf",
            "/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttf",
            "/usr/share/fonts/truetype/noto/NotoSansCJKsc-Regular.ttf",
            "/usr/share/fonts/truetype/wqy/wqy-microhei.ttc",
            "/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc",
            "/usr/local/share/fonts/NotoSansCJK-Regular.ttc",
        ):
            yield Path(p)


class ExportService:
    """导出服务"""

    def __init__(self, novel_repository: NovelRepository, chapter_repository: ChapterRepository):
        self.novel_repository = novel_repository
        self.chapter_repository = chapter_repository
        self._story_node_repo = None

    def _get_story_node_repo(self):
        if self._story_node_repo is None:
            try:
                from infrastructure.persistence.database.story_node_repository import StoryNodeRepository
                from application.paths import DATA_DIR
                self._story_node_repo = StoryNodeRepository(str(DATA_DIR / "aitext.db"))
            except Exception as e:
                logger.debug("无法加载 StoryNodeRepository: %s", e)
        return self._story_node_repo

    def _get_structure_tree(self, novel_id: str):
        """获取小说的结构树（部、卷、幕、章层级）"""
        repo = self._get_story_node_repo()
        if not repo:
            return None
        try:
            tree = repo.get_tree_sync(novel_id)
            return tree.to_tree_dict()
        except Exception as e:
            logger.debug("获取结构树失败: %s", e)
            return None

    def export_novel(self, novel_id: str, format: str) -> Tuple[bytes, str, str]:
        try:
            logger.info("开始导出小说: %s, 格式: %s", novel_id, format)
            novel = self.novel_repository.get_by_id(NovelId(novel_id))
            if not novel:
                raise ValueError(f"小说不存在: {novel_id}")
            chapters = self.chapter_repository.list_by_novel(NovelId(novel_id))
            chapters.sort(key=lambda x: x.number)
            
            structure_tree = self._get_structure_tree(novel_id)
            logger.info("导出: %s, 章节数 %s", novel.title, len(chapters))
            
            if format == "epub":
                result = self._export_to_epub(novel, chapters)
            elif format == "pdf":
                result = self._export_to_pdf(novel, chapters, structure_tree)
            elif format == "docx":
                result = self._export_to_docx(novel, chapters, structure_tree)
            elif format == "markdown":
                result = self._export_to_markdown(novel, chapters, structure_tree)
            else:
                raise ValueError(f"不支持的导出格式: {format}")
            logger.info("导出成功，%s 字节", len(result[0]))
            return result
        except ValueError:
            raise
        except Exception as e:
            logger.error("导出小说失败: %s", e, exc_info=True)
            raise

    def export_chapter(self, chapter_id: str, format: str) -> Tuple[bytes, str, str]:
        try:
            logger.info("开始导出章节: %s, 格式: %s", chapter_id, format)
            chapter = self.chapter_repository.get_by_id(ChapterId(chapter_id))
            if not chapter:
                raise ValueError(f"章节不存在: {chapter_id}")
            novel_id = chapter.novel_id.value if hasattr(chapter.novel_id, "value") else chapter.novel_id
            novel = self.novel_repository.get_by_id(NovelId(novel_id))
            if not novel:
                raise ValueError(f"小说不存在: {novel_id}")
            if format == "epub":
                result = self._export_to_epub(novel, [chapter])
            elif format == "pdf":
                result = self._export_to_pdf(novel, [chapter])
            elif format == "docx":
                result = self._export_to_docx(novel, [chapter])
            elif format == "markdown":
                result = self._export_to_markdown(novel, [chapter])
            else:
                raise ValueError(f"不支持的导出格式: {format}")
            data, mime, _ = result
            ext = {"epub": "epub", "pdf": "pdf", "docx": "docx", "markdown": "md"}[format]
            chapter_stem = _safe_filename_stem(
                f"{novel.title or 'novel'}-第{chapter.number}章"
            )
            logger.info("导出成功，%s 字节", len(data))
            return data, mime, f"{chapter_stem}.{ext}"
        except ValueError:
            raise
        except Exception as e:
            logger.error("导出章节失败: %s", e, exc_info=True)
            raise

    def _export_to_epub(self, novel: Novel, chapters: list[Chapter]) -> Tuple[bytes, str, str]:
        from ebooklib import epub

        book = epub.EpubBook()
        uid = _novel_id_str(novel)
        book.set_identifier(f"plotpilot:{uid}")
        book.set_title(novel.title or "未命名")
        book.set_language("zh")
        book.add_author(novel.author or "未知作者")

        intro = epub.EpubHtml(
            title="简介",
            file_name="intro.xhtml",
            lang="zh",
        )
        premise = html.escape((novel.premise or "").strip() or "（无简介）")
        intro.content = f"""<?xml version="1.0" encoding="utf-8"?>
<!DOCTYPE html>
<html xmlns="http://www.w3.org/1999/xhtml" lang="zh">
<head><title>简介</title><meta charset="utf-8"/></head>
<body>
<h1>{html.escape(novel.title or "未命名")}</h1>
<p>作者：{html.escape(novel.author or "—")}</p>
<p>{premise}</p>
</body>
</html>"""
        book.add_item(intro)

        spine_items: List[epub.EpubHtml] = [intro]
        for i, ch in enumerate(chapters):
            fname = f"chap_{i + 1:03d}.xhtml"
            title_txt = _chapter_display_title(ch)
            title_esc = html.escape(title_txt)
            body = _content_to_html_paragraphs(ch.content or "")
            item = epub.EpubHtml(title=title_txt, file_name=fname, lang="zh")
            item.content = f"""<?xml version="1.0" encoding="utf-8"?>
<!DOCTYPE html>
<html xmlns="http://www.w3.org/1999/xhtml" xmlns:epub="http://www.idpf.org/2007/ops" lang="zh">
<head><title>{title_esc}</title><meta charset="utf-8"/></head>
<body>
<h1>{title_esc}</h1>
{body}
</body>
</html>"""
            book.add_item(item)
            spine_items.append(item)

        book.toc = tuple([intro] + spine_items[1:])
        book.add_item(epub.EpubNcx())
        # 不使用空 EpubNav（ebooklib 生成 nav 时会解析正文，空文档会触发 lxml Document is empty）
        book.spine = spine_items

        tmp_path: Optional[str] = None
        try:
            fd, tmp_path = tempfile.mkstemp(suffix=".epub")
            os.close(fd)
            epub.write_epub(tmp_path, book, {})
            with open(tmp_path, "rb") as f:
                data = f.read()
        finally:
            if tmp_path and os.path.isfile(tmp_path):
                try:
                    os.unlink(tmp_path)
                except OSError:
                    pass

        stem = _safe_filename_stem(novel.title)
        return data, "application/epub+zip", f"{stem}.epub"

    def _try_register_cjk_font(self, pdf) -> bool:
        for path in _cjk_font_paths():
            if not path.is_file():
                continue
            try:
                pdf.add_font("PlotExportCJK", "", str(path), uni=True)
                return True
            except Exception as e:
                logger.debug("PDF 跳过字体 %s: %s", path, e)
        return False

    def _extract_structure_nodes(self, structure_tree) -> list[dict]:
        if not isinstance(structure_tree, dict):
            return []
        nodes = structure_tree.get("nodes")
        if isinstance(nodes, list):
            return nodes
        tree = structure_tree.get("tree")
        if isinstance(tree, dict) and isinstance(tree.get("nodes"), list):
            return tree["nodes"]
        return []

    def _structure_node_title(self, node: dict) -> str:
        node_type = str(node.get("node_type") or "").lower()
        title = str(node.get("title") or node.get("label") or "").strip()
        number = node.get("number")

        if title:
            compact = re.sub(r"\s+", "", title)
            if node_type == "part" and compact.startswith("第") and "部" in compact:
                return title
            if node_type == "volume" and compact.startswith("第") and "卷" in compact:
                return title
            if node_type == "act" and ("幕" in compact[:4] or compact.startswith("第")):
                return title
            if node_type == "chapter" and re.match(r"^第[0-9一二三四五六七八九十百千零两]+章", compact):
                return title

        if number is None:
            return title or "未命名"

        prefix_map = {
            "part": f"第{number}部",
            "volume": f"第{number}卷",
            "act": f"幕{number}",
            "chapter": f"第{number}章",
        }
        prefix = prefix_map.get(node_type, str(number))
        if title:
            return f"{prefix}：{title}"
        return prefix

    def _build_chapter_structure_map(self, structure_tree) -> dict[str, list[tuple[str, str]]]:
        chapter_map: dict[str, list[tuple[str, str]]] = {}
        nodes = self._extract_structure_nodes(structure_tree)
        if not nodes:
            return chapter_map

        def walk(items, ancestors):
            for node in items:
                node_type = str(node.get("node_type") or "").lower()
                current_ancestors = ancestors
                if node_type in {"part", "volume", "act"}:
                    current_ancestors = ancestors + [
                        (node_type, self._structure_node_title(node))
                    ]
                elif node_type == "chapter":
                    node_id = node.get("id")
                    number = node.get("number")
                    if node_id:
                        chapter_map[str(node_id)] = list(ancestors)
                    if number is not None:
                        chapter_map[f"number:{number}"] = list(ancestors)

                children = node.get("children") or []
                if children:
                    walk(children, current_ancestors)

        walk(nodes, [])
        return chapter_map

    def _export_to_pdf(self, novel: Novel, chapters: list[Chapter], structure_tree=None) -> Tuple[bytes, str, str]:
        from fpdf import FPDF

        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=14)
        pdf.set_left_margin(20)
        pdf.set_right_margin(20)
        
        font = "Helvetica"
        font_size = 11
        if self._try_register_cjk_font(pdf):
            font = "PlotExportCJK"
            font_size = 12

        def add_text(size: float, text: str, line_h: float, align: str = "") -> None:
            pdf.set_font(font, size=size)
            body = (text or "").strip() or " "
            try:
                pdf.multi_cell(0, line_h, body, new_x="LMARGIN", new_y="NEXT", align=align)
            except Exception as e:
                logger.warning("PDF multi_cell 回退: %s", e)
                pdf.set_font("Helvetica", size=size)
                safe = (text or "").encode("ascii", errors="replace").decode("ascii")
                pdf.multi_cell(0, line_h, safe or " ", new_x="LMARGIN", new_y="NEXT")

        def add_title(size: float, text: str, line_h: float) -> None:
            pdf.set_font(font, size=size)
            body = (text or "").strip() or " "
            pdf.cell(0, line_h, body, new_x="LMARGIN", new_y="NEXT", align="C")

        def add_bold_text(size: float, text: str, line_h: float, align: str = "") -> None:
            if font == "Helvetica":
                pdf.set_font(font, size=size, style="B")
            else:
                pdf.set_font(font, size=size)
            body = (text or "").strip() or " "
            try:
                pdf.multi_cell(0, line_h, body, new_x="LMARGIN", new_y="NEXT", align=align)
            except Exception as e:
                logger.warning("PDF bold multi_cell 回退: %s", e)
                pdf.set_font("Helvetica", size=size, style="B")
                safe = (text or "").encode("ascii", errors="replace").decode("ascii")
                pdf.multi_cell(0, line_h, safe or " ", new_x="LMARGIN", new_y="NEXT", align=align)

        pdf.add_page()
        
        add_title(20, novel.title or "未命名", 12)
        pdf.ln(6)
        
        author_line = f"作者：{novel.author or '未知作者'}"
        add_text(12, author_line, 7, align="C")
        pdf.ln(4)
        
        premise = (novel.premise or "").strip()
        if premise:
            add_bold_text(12, "简介", 7)
            add_text(font_size, premise, 6)
            pdf.ln(8)
        
        pdf.ln(8)

        chapter_structure_map = self._build_chapter_structure_map(structure_tree)
        last_ancestors: list[tuple[str, str]] = []
        heading_sizes = {
            "part": 16,
            "volume": 15,
            "act": 14,
        }

        for i, ch in enumerate(chapters):
            ancestors = (
                chapter_structure_map.get(str(ch.id))
                or chapter_structure_map.get(f"number:{ch.number}")
                or []
            )
            shared = 0
            max_shared = min(len(last_ancestors), len(ancestors))
            while shared < max_shared and last_ancestors[shared] == ancestors[shared]:
                shared += 1

            if i > 0:
                pdf.ln(4)

            for node_type, heading in ancestors[shared:]:
                add_bold_text(heading_sizes.get(node_type, 13), heading, 8)
                pdf.ln(1)

            add_bold_text(13, _chapter_display_title(ch), 8)
            pdf.ln(4)
            
            content = (ch.content or "").strip()
            if content:
                lines = content.split("\n")
                for line in lines:
                    line = line.strip()
                    if line:
                        add_text(font_size, line, 6)
                    else:
                        pdf.ln(3)
            else:
                add_text(font_size, "（无正文）", 6)

            last_ancestors = list(ancestors)

        out = pdf.output()
        if isinstance(out, str):
            data = out.encode("latin-1")
        elif isinstance(out, bytearray):
            data = bytes(out)
        else:
            data = bytes(out) if hasattr(out, '__bytes__') else out
        
        stem = _safe_filename_stem(novel.title)
        return data, "application/pdf", f"{stem}.pdf"

    def _export_to_docx(self, novel: Novel, chapters: list[Chapter], structure_tree=None) -> Tuple[bytes, str, str]:
        from docx import Document

        doc = Document()
        doc.add_heading(novel.title or "未命名", level=0)
        doc.add_paragraph(f"作者：{novel.author or '—'}")
        p_pre = doc.add_paragraph()
        p_pre.add_run("简介：").bold = True
        p_pre.add_run((novel.premise or "").strip() or "（无）")

        chapter_structure_map = self._build_chapter_structure_map(structure_tree)
        last_ancestors: list[tuple[str, str]] = []
        heading_levels = {
            "part": 1,
            "volume": 2,
            "act": 3,
        }

        for ch in chapters:
            ancestors = (
                chapter_structure_map.get(str(ch.id))
                or chapter_structure_map.get(f"number:{ch.number}")
                or []
            )
            shared = 0
            max_shared = min(len(last_ancestors), len(ancestors))
            while shared < max_shared and last_ancestors[shared] == ancestors[shared]:
                shared += 1

            for node_type, heading in ancestors[shared:]:
                doc.add_heading(heading, level=heading_levels.get(node_type, 3))

            doc.add_heading(_chapter_display_title(ch), level=4)
            content = ch.content or ""
            if not content.strip():
                doc.add_paragraph("（无正文）")
                last_ancestors = list(ancestors)
                continue
            for line in content.splitlines():
                doc.add_paragraph(line)
            last_ancestors = list(ancestors)

        buf = io.BytesIO()
        doc.save(buf)
        stem = _safe_filename_stem(novel.title)
        return (
            buf.getvalue(),
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            f"{stem}.docx",
        )

    def _render_docx_structure(self, doc, nodes, indent=0):
        """递归渲染结构树到docx"""
        from docx.shared import Pt
        
        for node in nodes:
            node_type = node.get("node_type")
            title = node.get("title", "") or node.get("label", "")
            
            if title:
                text = " " * indent + title
            else:
                number = node.get("number")
                if number:
                    if node_type == "part":
                        text = " " * indent + f"第{number}部"
                    elif node_type == "volume":
                        text = " " * indent + f"第{number}卷"
                    elif node_type == "act":
                        text = " " * indent + f"幕{number}"
                    elif node_type == "chapter":
                        text = " " * indent + f"第{number}章"
                    else:
                        text = " " * indent + str(number)
                else:
                    text = " " * indent + "未命名"
            
            p = doc.add_paragraph(text)
            p.paragraph_format.left_indent = Pt(indent * 2)
            
            children = node.get("children", [])
            if children:
                self._render_docx_structure(doc, children, indent + 4)

    def _export_to_markdown(self, novel: Novel, chapters: list[Chapter], structure_tree=None) -> Tuple[bytes, str, str]:
        lines: List[str] = [
            f"# {novel.title or '未命名'}",
            "",
            f"**作者**: {novel.author or '—'}",
            "",
            "## 简介",
            "",
            (novel.premise or "").strip() or "（无）",
            "",
        ]

        chapter_structure_map = self._build_chapter_structure_map(structure_tree)
        last_ancestors: list[tuple[str, str]] = []
        heading_prefix = {
            "part": "##",
            "volume": "###",
            "act": "####",
        }

        for ch in chapters:
            ancestors = (
                chapter_structure_map.get(str(ch.id))
                or chapter_structure_map.get(f"number:{ch.number}")
                or []
            )
            shared = 0
            max_shared = min(len(last_ancestors), len(ancestors))
            while shared < max_shared and last_ancestors[shared] == ancestors[shared]:
                shared += 1

            for node_type, heading in ancestors[shared:]:
                lines.append(f"{heading_prefix.get(node_type, '####')} {heading}")
                lines.append("")

            lines.append(f"##### {_chapter_display_title(ch)}")
            lines.append("")
            lines.append((ch.content or "").strip() or "（无正文）")
            lines.append("")
            last_ancestors = list(ancestors)
        text = "\n".join(lines)
        stem = _safe_filename_stem(novel.title)
        return text.encode("utf-8"), "text/markdown; charset=utf-8", f"{stem}.md"

    def _render_markdown_structure(self, lines, nodes, indent=0):
        """递归渲染结构树到markdown"""
        for node in nodes:
            node_type = node.get("node_type")
            title = node.get("title", "") or node.get("label", "")
            
            if title:
                text = " " * indent + title
            else:
                number = node.get("number")
                if number:
                    if node_type == "part":
                        text = " " * indent + f"第{number}部"
                    elif node_type == "volume":
                        text = " " * indent + f"第{number}卷"
                    elif node_type == "act":
                        text = " " * indent + f"幕{number}"
                    elif node_type == "chapter":
                        text = " " * indent + f"第{number}章"
                    else:
                        text = " " * indent + str(number)
                else:
                    text = " " * indent + "未命名"
            
            lines.append(text)
            
            children = node.get("children", [])
            if children:
                self._render_markdown_structure(lines, children, indent + 4)

    def export_library(self, format: str) -> Tuple[bytes, str, str]:
        """导出书目（所有小说列表）
        
        Args:
            format: 导出格式，支持 pdf, markdown
        
        Returns:
            元组：(文件内容, MIME类型, 文件名)
        """
        try:
            logger.info("开始导出书目, 格式: %s", format)
            novels = self.novel_repository.list_all()
            novels.sort(key=lambda x: x.created_at or 0)
            logger.info("书目导出: %d 本小说", len(novels))
            
            if format == "pdf":
                result = self._export_library_to_pdf(novels)
            elif format == "markdown":
                result = self._export_library_to_markdown(novels)
            else:
                raise ValueError(f"不支持的书目导出格式: {format}")
            
            logger.info("书目导出成功，%d 字节", len(result[0]))
            return result
        except ValueError:
            raise
        except Exception as e:
            logger.error("书目导出失败: %s", e, exc_info=True)
            raise

    def _export_library_to_pdf(self, novels: list[Novel]) -> Tuple[bytes, str, str]:
        from fpdf import FPDF

        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=14)
        pdf.set_left_margin(20)
        pdf.set_right_margin(20)
        
        font = "Helvetica"
        font_size = 11
        if self._try_register_cjk_font(pdf):
            font = "PlotExportCJK"
            font_size = 12

        def add_text(size: float, text: str, line_h: float, align: str = "") -> None:
            pdf.set_font(font, size=size)
            body = (text or "").strip() or " "
            try:
                pdf.multi_cell(0, line_h, body, new_x="LMARGIN", new_y="NEXT", align=align)
            except Exception as e:
                logger.warning("PDF multi_cell 回退: %s", e)
                pdf.set_font("Helvetica", size=size)
                safe = (text or "").encode("ascii", errors="replace").decode("ascii")
                pdf.multi_cell(0, line_h, safe or " ", new_x="LMARGIN", new_y="NEXT")

        def add_bold_text(size: float, text: str, line_h: float, align: str = "") -> None:
            if font == "Helvetica":
                pdf.set_font(font, size=size, style="B")
            else:
                pdf.set_font(font, size=size)
            body = (text or "").strip() or " "
            pdf.cell(0, line_h, body, new_x="LMARGIN", new_y="NEXT", align=align)

        pdf.add_page()
        
        add_bold_text(20, "我的书目", 12, align="C")
        pdf.ln(6)
        
        from datetime import datetime
        now = datetime.now().strftime("%Y年%m月%d日 %H:%M")
        add_text(10, f"生成时间：{now}", 5, align="C")
        pdf.ln(10)

        add_bold_text(12, f"共 {len(novels)} 本小说", 7)
        pdf.ln(4)

        for i, novel in enumerate(novels, 1):
            add_bold_text(font_size + 1, f"{i}. {novel.title or '未命名'}", 8)
            
            info_lines = []
            if novel.author:
                info_lines.append(f"作者：{novel.author}")
            if novel.target_chapters:
                info_lines.append(f"目标章节：{novel.target_chapters} 章")
            if novel.stage:
                stage_map = {
                    "planning": "规划中",
                    "writing": "创作中",
                    "completed": "已完成"
                }
                info_lines.append(f"状态：{stage_map.get(novel.stage.value, novel.stage.value)}")
            
            if info_lines:
                add_text(font_size - 1, " / ".join(info_lines), 5)
            
            premise = (novel.premise or "").strip()
            if premise:
                pdf.ln(2)
                if len(premise) > 150:
                    premise = premise[:150] + "..."
                add_text(font_size - 1, f"简介：{premise}", 5)
            
            pdf.ln(8)

        out = pdf.output()
        if isinstance(out, str):
            data = out.encode("latin-1")
        elif isinstance(out, bytearray):
            data = bytes(out)
        else:
            data = bytes(out) if hasattr(out, '__bytes__') else out
        
        return data, "application/pdf", "我的书目.pdf"

    def _export_library_to_markdown(self, novels: list[Novel]) -> Tuple[bytes, str, str]:
        from datetime import datetime
        
        lines: List[str] = [
            "# 我的书目",
            "",
            f"生成时间：{datetime.now().strftime('%Y年%m月%d日 %H:%M')}",
            "",
            f"共 {len(novels)} 本小说",
            "",
        ]
        
        stage_map = {
            "planning": "规划中",
            "writing": "创作中",
            "completed": "已完成"
        }
        
        for i, novel in enumerate(novels, 1):
            lines.append(f"## {i}. {novel.title or '未命名'}")
            lines.append("")
            
            info_parts = []
            if novel.author:
                info_parts.append(f"**作者**：{novel.author}")
            if novel.target_chapters:
                info_parts.append(f"**目标章节**：{novel.target_chapters} 章")
            if novel.stage:
                info_parts.append(f"**状态**：{stage_map.get(novel.stage.value, novel.stage.value)}")
            
            if info_parts:
                lines.append(" | ".join(info_parts))
                lines.append("")
            
            premise = (novel.premise or "").strip()
            if premise:
                lines.append("**简介**：")
                lines.append("> " + premise)
                lines.append("")
        
        text = "\n".join(lines)
        return text.encode("utf-8"), "text/markdown; charset=utf-8", "我的书目.md"
